from flask import Flask, render_template, redirect, request, send_file
from docx import Document

app = Flask(__name__)

# Datos ficticios de usuarios
usuarios = {
    'ejecutivo': {'username': 'ejecutivo', 'password': '1234'},
    'Adm': {'username': 'Adm', 'password': '123abc'},
    'Rbarrios': {'username': 'Rbarrios', 'password': '71386286'},
    'Mtarazona': {'username': 'Mtarazona', 'password': '1234'},
    'Sandy': {'username': 'Sgarcia', 'password': 'abc0309'},
    'trabajador': {'username': 'trabajador', 'password': 'abcd'}
}

# Rutas
@app.route('/')
def inicio():
    return render_template('login.html')

@app.route('/login', methods=['POST'])
def login():
    username = request.form['username']
    password = request.form['password']

    if username in usuarios and usuarios[username]['password'] == password:
        if username == 'ejecutivo':
            return redirect('/perfil_ejecutivo')
        if username == 'Adm':
            return redirect('/perfil_ejecutivo')
        if username == 'Rbarrios':
            return redirect('perfil_trabajador')
        if username == 'Sgarcia':
            return redirect('perfil_trabajador')
        if username == 'Mtarazona':
            return redirect('perfil_trabajador')
        elif username == 'trabajador':
            return redirect('/perfil_trabajador')
    else:
        return "Credenciales incorrectas. Por favor, inténtalo de nuevo."

@app.route('/perfil_ejecutivo')
def perfil_ejecutivo():
    return render_template('formulario_registro.html', perfil='Ejecutivo')

@app.route('/perfil_trabajador')
def perfil_trabajador():
    return render_template('formulario_registro.html', perfil='Trabajador')

@app.route('/guardar_registro', methods=['POST'])
def guardar_registro():
    # Obtener los datos del formulario
    nombre = request.form['nombre']
    apellido = request.form['apellido']
    edad = request.form['edad']
    genero = request.form['genero']
    direccion = request.form['direccion']

    # Crear un documento Word y agregar los datos
    document = Document()
    document.add_heading('Registro de Datos', level=1)
    document.add_paragraph(f'Nombre: {nombre}')
    document.add_paragraph(f'Apellido: {apellido}')
    document.add_paragraph(f'Edad: {edad}')
    document.add_paragraph(f'Género: {genero}')
    document.add_paragraph(f'Dirección: {direccion}')

    # Guardar el documento en memoria
    docx_file = 'registro.docx'
    document.save(docx_file)

    # Descargar el archivo Word
    return send_file(docx_file, as_attachment=True)


if __name__ == "__main__":
    app.run(debug=True, host='0.0.0.0', port=5004)  # Cambia 5002 por el puerto que desees usar
    
